# 文档解析器
"""
支持多种文档格式的解析:
- PDF
- DOCX
- TXT
"""

from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

from src.config import get_settings
from src.utils import get_logger

logger = get_logger(__name__)


class ParsedDocument(BaseModel):
    """解析后的文档"""
    filename: str
    file_type: str
    text: str
    pages: int = 1
    metadata: dict[str, Any] = Field(default_factory=dict)
    tables: list[dict] = Field(default_factory=list)
    images: list[dict] = Field(default_factory=list)
    error: str | None = None


class DocumentParser:
    """文档解析器
    
    支持:
    - PDF 解析 (使用 PyMuPDF)
    - DOCX 解析
    - TXT 解析
    - 表格提取
    - 图片提取
    """
    
    def __init__(self):
        settings = get_settings()
        parser_config = settings.ingestion.parser
        
        self.supported_formats = parser_config.supported_formats
        self.max_file_size = parser_config.max_file_size
    
    def parse(self, file_path: str | Path) -> ParsedDocument:
        """解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            ParsedDocument: 解析结果
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return ParsedDocument(
                filename=file_path.name,
                file_type="unknown",
                text="",
                error=f"File not found: {file_path}",
            )
        
        # 检查文件大小
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            return ParsedDocument(
                filename=file_path.name,
                file_type=file_path.suffix,
                text="",
                error=f"File too large: {file_size} bytes (max: {self.max_file_size})",
            )
        
        # 检查文件类型
        suffix = file_path.suffix.lower()
        if suffix not in self.supported_formats:
            return ParsedDocument(
                filename=file_path.name,
                file_type=suffix,
                text="",
                error=f"Unsupported file type: {suffix}",
            )
        
        # 根据类型解析
        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix == ".docx":
            return self._parse_docx(file_path)
        elif suffix == ".txt":
            return self._parse_txt(file_path)
        else:
            return ParsedDocument(
                filename=file_path.name,
                file_type=suffix,
                text="",
                error=f"Parser not implemented for: {suffix}",
            )
    
    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """解析 PDF 文档"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            text_parts = []
            tables = []
            images = []
            
            for page_num, page in enumerate(doc):
                # 提取文本
                text = page.get_text()
                text_parts.append(text)
                
                # 提取表格 (简单实现)
                # TODO: 使用更高级的表格提取库
                
                # 提取图片信息
                image_list = page.get_images()
                for img in image_list:
                    images.append({
                        "page": page_num + 1,
                        "xref": img[0],
                        "width": img[2],
                        "height": img[3],
                    })
            
            metadata = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "keywords": doc.metadata.get("keywords", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
            }
            
            doc.close()
            
            return ParsedDocument(
                filename=file_path.name,
                file_type=".pdf",
                text="\n\n".join(text_parts),
                pages=len(text_parts),
                metadata=metadata,
                tables=tables,
                images=images,
            )
            
        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install pymupdf")
            return ParsedDocument(
                filename=file_path.name,
                file_type=".pdf",
                text="",
                error="PyMuPDF not installed",
            )
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return ParsedDocument(
                filename=file_path.name,
                file_type=".pdf",
                text="",
                error=str(e),
            )
    
    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """解析 DOCX 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            text_parts = []
            tables = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({"data": table_data})
            
            # 提取元数据
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "keywords": doc.core_properties.keywords or "",
            }
            
            return ParsedDocument(
                filename=file_path.name,
                file_type=".docx",
                text="\n\n".join(text_parts),
                pages=1,  # DOCX 不分页
                metadata=metadata,
                tables=tables,
            )
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ParsedDocument(
                filename=file_path.name,
                file_type=".docx",
                text="",
                error="python-docx not installed",
            )
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return ParsedDocument(
                filename=file_path.name,
                file_type=".docx",
                text="",
                error=str(e),
            )
    
    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        """解析 TXT 文档"""
        try:
            # 尝试多种编码
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            return ParsedDocument(
                filename=file_path.name,
                file_type=".txt",
                text=text,
                pages=1,
            )
            
        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            return ParsedDocument(
                filename=file_path.name,
                file_type=".txt",
                text="",
                error=str(e),
            )
    
    async def parse_async(self, file_path: str | Path) -> ParsedDocument:
        """异步解析文档（实际上是同步，但提供异步接口）"""
        import asyncio
        return await asyncio.get_event_loop().run_in_executor(
            None, self.parse, file_path
        )

